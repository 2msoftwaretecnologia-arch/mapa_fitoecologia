from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def gerar_word(texto):
    """
    Gera um arquivo Word com texto imputável.
    Define a borda com base no número de caracteres do texto.
    O texto será justificado e em Times New Roman, 11 pt.
    """
    n = len(texto)

    # Define intervalos baseados no tamanho do texto
    if n > 0 and n <= 340:
        borda = 6.25
    elif n > 340 and n <= 490:
        borda = 5.75
    elif n > 490 and n <= 800:
        borda = 8
    elif n > 800:
        borda = 6.5
    else:
        borda = 6  # valor padrão

    # Criar documento
    doc = Document()

    # Aplicar margens
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(borda)
    section.right_margin = Cm(borda)

    # Adicionar parágrafo justificado
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run(texto)
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Salvar
    doc.save("saida.docx")
    print(f"✅ Arquivo 'saida.docx' gerado com sucesso! (margem = {borda} cm, fonte Times New Roman 7pt)")

# ==============================
# Execução do programa
# ==============================
texto_input = str(input("Digite o texto para gerar o Word: "))
gerar_word(texto_input)
